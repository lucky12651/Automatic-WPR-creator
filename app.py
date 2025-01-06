from datetime import datetime, timedelta
from docx import Document
import openai
from docx.shared import Pt


# Set your OpenAI API Key
openai.api_key = ""

# Function to generate valid dates for the WPR
def generate_wpr_dates(start_date_str, week_number):
    start_date = datetime.strptime(start_date_str, "%d/%m/%Y")
    start_of_week = start_date + timedelta(weeks=week_number - 1)
    end_of_week = start_of_week + timedelta(days=6)
    
    # Format dates into the required string format
    start_date_str = start_of_week.strftime("%d/%m/%Y")
    end_date_str = end_of_week.strftime("%d/%m/%Y")
    
    return start_date_str, end_date_str

# Function to calculate remaining WPRs
def calculate_remaining_wprs(total_wprs, current_week):
    remaining_wprs = total_wprs - current_week
    return remaining_wprs


# Function to generate content using ChatGPT for each section, ensuring uniqueness for every week
def generate_section_content(section_title, base_content, week_number):
    prompt = f"""
    Based on the following Weekly Progress Report, create unique and new content for the section "{section_title}" for Week {week_number}. Ensure that this new content is different from the previous weeks' content, incorporating any updates or improvements:
    {base_content}
    """
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are an assistant generating weekly progress report content, ensuring the content is unique and relevant to the specified week."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=500,
        temperature=0.7,
    )
    return response['choices'][0]['message']['content'].strip()

# Function to replace placeholders in the Word template and format text
def replace_placeholders(doc, content_replacements):
    for paragraph in doc.paragraphs:
        for placeholder, content in content_replacements.items():
            if placeholder in paragraph.text:
                
                # Clean up Markdown-like bold formatting (removing **)
                content = content.replace("**", "")
                content = content.replace("###", "")
                content = content.replace("##", "")
                # Replace the placeholder text
                paragraph.text = paragraph.text.replace(placeholder, content)
                # Change font size to 11 and remove bold formatting
                for run in paragraph.runs:
                    run.font.size = Pt(11)  # Set font size to 11
                    run.font.bold = False  # Ensure text is not bold

# Function to update table content in the Word document
def update_table_content(doc, table_data):
    for table in doc.tables:
        if "Days/Time" in table.rows[0].cells[0].text:  # Ensure correct table
            for i, row_data in enumerate(table_data):
                if i == 0:
                    continue  # Skip headers
                # Add a new row for the table
                row = table.add_row()
                # Iterate over the cells and update the text
                for j, cell_data in enumerate(row_data):
                    # Clean up the content to remove '**' or any unwanted Markdown-like symbols
                    clean_cell_data = cell_data.replace("**", "").replace("###", "").replace("##", "")
                    row.cells[j].text = clean_cell_data
            break  # Ensure only the first table is updated


# Main function to generate WPRs
def generate_wprs(base_content, template_path, start_week, total_weeks):
    for week in range(start_week, start_week + total_weeks):
        # Generate unique content for each section
        targets_set = generate_section_content("TARGETS SET FOR THE WEEK", base_content, week)
        achievements = generate_section_content("ACHIEVEMENTS FOR THE WEEK", base_content, week)
        future_work = generate_section_content("FUTURE WORK PLANS", base_content, week)
        summary = generate_section_content("WEEK'S SUMMARY", base_content, week)

        # Generate date and WPR text using the new date generation function
        start_date, end_date = generate_wpr_dates("17/12/2024", week)  # Modify the starting date as per your pattern
        date_text = f"{start_date}-{end_date}"
        wpr_text = f"{week}"

        # Calculate remaining WPRs
        total_wprs = 18  # Assuming total WPRs are 18
        remaining_wprs = calculate_remaining_wprs(total_wprs, week)

        # Generate table data dynamically
        table_data = [
            ["Days/Time", "Tasks"],
            ["Monday", generate_section_content("Monday Activities", base_content, week)],
            ["Tuesday", generate_section_content("Tuesday Activities", base_content, week)],
            ["Wednesday", generate_section_content("Wednesday Activities", base_content, week)],
            ["Thursday", generate_section_content("Thursday Activities", base_content, week)],
            ["Friday", generate_section_content("Friday Activities", base_content, week)],
        ]

        # Define replacements for text placeholders
        content_replacements = {
            "{{TARGETS_SET}}": targets_set,
            "{{ACHIEVEMENTS}}": achievements,
            "{{FUTURE_WORK}}": future_work,
            "{{SUMMARY}}": summary,
            "{{DATE_TEXT}}": date_text,
            "{{WPR_TEXT}}": wpr_text,
            "{{WPR_REMAINING}}": str(remaining_wprs),
        }

        # Load the Word template
        doc = Document(template_path)

        # Replace text placeholders and format content
        replace_placeholders(doc, content_replacements)

        # Update table content
        update_table_content(doc, table_data)

        # Save the updated document
        output_path = f"WPR_Week_{week}.docx"
        doc.save(output_path)
        print(f"WPR for Week {week} saved as {output_path}")

# Base WPR content


# Function to update the WPR from the txt file
def read_base_wpr_content(file_path):
    with open(file_path, 'r') as file:
        return file.read()

# Path to your base content text file
base_wpr_file_path = "hello.txt"

# Read the base WPR content from the txt file
base_wpr_content = read_base_wpr_content(base_wpr_file_path)



# Path to your Word template
template_path = "WPR.docx"

# Generate 2 WPRs starting from Week 4
generate_wprs(base_wpr_content, template_path, start_week=4, total_weeks=18)
